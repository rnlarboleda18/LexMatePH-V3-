import os
import re
from pathlib import Path
import fitz  # PyMuPDF
from docx import Document

ROC_DIR = Path(r"C:\Users\rnlar\.gemini\antigravity\scratch\bar_project_v2\LexCode\Codals\md\ROC")

def convert_docx_to_md(docx_path):
    """Convert .docx to Markdown with run-level styling."""
    doc = Document(docx_path)
    md_lines = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # Basic Heading Detection based on style name
        style_name = para.style.name.lower()
        if 'heading 1' in style_name:
            md_lines.append(f"# {text}\n")
            continue
        elif 'heading 2' in style_name:
            md_lines.append(f"## {text}\n")
            continue
        elif 'heading 3' in style_name:
            md_lines.append(f"### {text}\n")
            continue
        elif 'heading' in style_name:
            md_lines.append(f"#### {text}\n")
            continue
            
        # Build paragraph text from runs to preserve inline styling
        para_text = ""
        for run in para.runs:
            run_text = run.text
            if not run_text:
                continue
            
            cleaned_run_text = run_text  # Avoid stripping run levels to maintain spaces
            if run.bold:
                 cleaned_run_text = f"**{cleaned_run_text}**"
            if run.italic:
                 cleaned_run_text = f"*{cleaned_run_text}*"
            if run.underline:
                 cleaned_run_text = f"<u>{cleaned_run_text}</u>"
            
            para_text += cleaned_run_text
            
        if para_text.strip():
            md_lines.append(f"{para_text.strip()}\n")
                
    # Handle Tables
    for table in doc.tables:
        md_lines.append("\n")
        for i, row in enumerate(table.rows):
            row_data = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            md_lines.append("| " + " | ".join(row_data) + " |")
            if i == 0:
                md_lines.append("|" + "---| " * len(row_data) + "|")
        md_lines.append("\n")
                
    return "\n".join(md_lines)

def convert_pdf_to_md(pdf_path):
    """Convert .pdf to Markdown using PyMuPDF dictionary extraction to find underline styling."""
    doc = fitz.open(pdf_path)
    md_lines = []
    TEXT_COLLECT_STYLES = 32768
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        d1 = page.get_text("dict", flags=TEXT_COLLECT_STYLES)
        
        # Pull vector drawings to find underlines misses
        drawings = page.get_drawings()
        underline_rects = []
        for d in drawings:
             if "items" in d:
                  for item in d["items"]:
                       if item[0] == "l": # line
                            # Create bounding box for line segment
                            r = fitz.Rect(item[1], item[2])
                            if abs(r.y1 - r.y0) < 3 and r.width > 2:
                                 underline_rects.append(r)
        
        for block in d1.get("blocks", []):
            if "lines" not in block:
                continue
                
            # Filter Running Headers/Footers (bounding box tests)
            if "bbox" in block:
                y0 = block["bbox"][1]
                y1 = block["bbox"][3]
                if y0 > 690:  # Footer zone
                    continue
                if y1 < 65:   # Header zone
                    continue
                
            block_text = ""
            for line in block["lines"]:
                line_text = ""
                for span in line.get("spans", []):
                    span_text = span.get("text", "")
                    if not span_text.strip():
                        line_text += span_text
                        continue
                        
                    line_text += span_text
                    
                if line_text.strip():
                     block_text += line_text + " "
                     
            cleaned_block_text = block_text.strip().replace("\n", " ")
            if not cleaned_block_text:
                continue
                
                
            # Basic Header detection
            if len(cleaned_block_text) < 100 and cleaned_block_text.isupper():
                if cleaned_block_text.startswith("RULE") or cleaned_block_text.startswith("SECTION"):
                     md_lines.append(f"### {cleaned_block_text}\n")
                else:
                     md_lines.append(f"## {cleaned_block_text}\n")
            else:
                 md_lines.append(f"{cleaned_block_text}\n")
        
    return "\n".join(md_lines)

def main():
    if not ROC_DIR.exists():
        print(f"Directory {ROC_DIR} does not exist.")
        return
        
    files = list(ROC_DIR.glob("*"))
    print(f"Scanning {len(files)} items in {ROC_DIR}...")
    
    for file_path in files:
        ext = file_path.suffix.lower()
        if ext == ".md":
             continue
             
        print(f"Processing: {file_path.name} ({ext})...")
        md_content = ""
        
        try:
            if ext == ".docx":
                md_content = convert_docx_to_md(file_path)
            elif ext == ".pdf":
                md_content = convert_pdf_to_md(file_path)
            else:
                print(f"Skipping unsupported type: {ext}")
                continue
                
            if md_content:
                 output_name = file_path.stem + ".md"
                 output_path = ROC_DIR / output_name
                 
                 with open(output_path, "w", encoding="utf-8") as f:
                     f.write(md_content)
                 print(f"Saved: {output_name}")
            else:
                 print(f"Warning: No content extracted for {file_path.name}")
                 
        except Exception as e:
            print(f"Error converting {file_path.name}: {e}")

if __name__ == "__main__":
    main()
