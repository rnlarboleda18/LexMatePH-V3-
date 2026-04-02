from docx import Document

def inspect_docx(file_path):
    doc = Document(file_path)
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    
    # Inspect first 100 paragraphs for styles and alignment
    for i, para in enumerate(doc.paragraphs[:100]):
        style = para.style.name
        text = para.text.strip()
        if not text:
            continue
            
        # Check for bold, alignment, or specific keywords
        is_bold = any(run.bold for run in para.runs)
        alignment = para.alignment
        
        print(f"P{i} [{style}] [Bold:{is_bold}] [Align:{alignment}]: {text[:80]}")

if __name__ == "__main__":
    inspect_docx(r"C:\Users\rnlar\.gemini\antigravity\scratch\bar_project_v2\LexCode\Codals\Word\1987 Philippine Constitution.docx")
